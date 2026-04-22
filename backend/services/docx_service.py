"""
docx_service.py
───────────────
Generates a styled DOCX assessment report using python-docx.
"""

import os
import uuid
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import TEMP_DIR
from schemas.models import Section


# ─────────────────────────────────────────────
# Brand colours (RGB tuples)
# ─────────────────────────────────────────────
RED         = RGBColor(0xC3, 0x1D, 0x27)
DARK_GRAY   = RGBColor(0x11, 0x18, 0x27)
MID_GRAY    = RGBColor(0x6B, 0x72, 0x80)
GREEN       = RGBColor(0x16, 0x65, 0x34)
LIGHT_GRAY  = RGBColor(0xF3, 0xF4, 0xF6)


def _set_cell_bg(cell, hex_color: str):
    """Set background colour of a table cell."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_run(para, text: str, bold=False, italic=False, color: RGBColor = None, size_pt: int = 10):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def generate_docx(sections: List[Section], jd_snippet: str = "") -> str:
    """
    Build a DOCX from assessment sections and return the temp file path.
    """
    file_path = os.path.join(TEMP_DIR, f"assessment_{uuid.uuid4().hex[:8]}.docx")
    doc = Document()

    # ── Page margins ──────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Document title ────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(title_para, "ICS AI Hiring Assessment", bold=True, color=DARK_GRAY, size_pt=20)

    sub_para = doc.add_paragraph()
    _add_run(sub_para, "Assessment Generator  |  AI-powered candidate screening", color=MID_GRAY, size_pt=9)

    # Horizontal rule via paragraph border
    def _add_hr(doc_obj, color_hex="C31D27"):
        p = doc_obj.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    _add_hr(doc)

    total_q = sum(len(s.questions) for s in sections)
    meta = doc.add_paragraph()
    _add_run(meta, f"Total Questions: ", bold=True, color=DARK_GRAY, size_pt=9)
    _add_run(meta, str(total_q), bold=True, color=RED, size_pt=9)
    _add_run(meta, f"   |   Sections: ", bold=True, color=DARK_GRAY, size_pt=9)
    _add_run(meta, str(len(sections)), bold=True, color=RED, size_pt=9)
    doc.add_paragraph()

    # ── Sections & Questions ──────────────────
    q_global = 1
    for section in sections:
        # Domain banner (table with 1 cell, red background)
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.rows[0].cells[0]
        _set_cell_bg(cell, "C31D27")
        cell.width = Inches(6.5)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _add_run(p, f"  {section.domain}", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size_pt=12)
        doc.add_paragraph()

        for q in section.questions:
            # Question number
            qn_para = doc.add_paragraph()
            _add_run(qn_para, f"Q{q_global}. ", bold=True, color=RED, size_pt=10)
            _add_run(qn_para, q.question, bold=True, color=DARK_GRAY, size_pt=10)

            # Options
            options = q.options.model_dump()
            for letter, text in options.items():
                is_correct = letter == q.correct_answer
                opt_para = doc.add_paragraph(style="List Bullet")
                opt_para.paragraph_format.left_indent = Cm(0.5)
                if is_correct:
                    _add_run(opt_para, f"({letter})  {text}  ✓", bold=True, color=GREEN, size_pt=9)
                else:
                    _add_run(opt_para, f"({letter})  {text}", color=DARK_GRAY, size_pt=9)

            # Explanation box (table)
            exp_tbl = doc.add_table(rows=1, cols=1)
            exp_tbl.style = "Table Grid"
            exp_cell = exp_tbl.rows[0].cells[0]
            _set_cell_bg(exp_cell, "F3F4F6")
            exp_p = exp_cell.paragraphs[0]
            _add_run(exp_p, "💡 ", bold=False, size_pt=9)
            _add_run(exp_p, q.explanation, italic=True, color=MID_GRAY, size_pt=9)

            doc.add_paragraph()
            q_global += 1

        doc.add_paragraph()

    # ── Footer ────────────────────────────────
    _add_hr(doc, "E5E7EB")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(footer_para, "Generated by ICS AI Assessment Generator", italic=True, color=MID_GRAY, size_pt=8)

    doc.save(file_path)
    return file_path
